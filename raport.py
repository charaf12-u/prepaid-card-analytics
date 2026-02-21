
from docx import Document
from docx.shared import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd


def create_report_word( agregation, analyse_jeur, analyse_Semaine, analyse_Mois, clients_10Tops, clients_Ayant_Non_Regles, top_resturants,
        top_heur, df_Anormalis, s_User, r, filename="documents_rapport/rapport.docx" 
    ):

    pd.set_option("display.max_columns", None)

    doc = Document()

    title = doc.add_heading("RAPPORT JOURNALIER", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_section(title, df):
        doc.add_heading(title, 3)

        p = doc.add_paragraph()
        run = p.add_run(df.head().to_string())
        run.font.name = "Courier New"
        run.font.size = Pt(9)

        info = doc.add_paragraph(f"Nombre de lignes = {len(df)}")
        info.runs[0].italic = True

        doc.add_paragraph("")

    add_section(
        "Montants totaux par période Jour / Semaine / Mois :",
        agregation.reset_index()
    )

    doc.add_heading("Moyenne des soldes des cartes prépayées :", 3).underline = True
    add_section("Jour :", analyse_jeur)
    doc.add_picture("documents_rapport/Moyenne_solde_par_jeur.jpg" , width=Inches(4))
    add_section("Semaine :", analyse_Semaine)
    doc.add_picture("documents_rapport/Moyenne_solde_par_Semaine.jpg" , width=Inches(4))
    add_section("Mois :", analyse_Mois)
    doc.add_picture("documents_rapport/Moyenne_solde_par_mois.jpg" , width=Inches(4))
    

    add_section("Clients les plus dépensiers :", clients_10Tops)
    add_section("Clients à risque :", clients_Ayant_Non_Regles.reset_index())

    add_section("Top restaurants de pointe :", top_resturants.reset_index())
    add_section("Top heures de pointe :", top_heur.reset_index())

    add_section("Anomalies détectées :", df_Anormalis.reset_index())
    add_section("Performance des Caissiers :", s_User.reset_index())

    doc.add_heading(
        "Corrélation entre colonne Solde_CPP et Montant_Rgl :", 3
    )
    doc.add_paragraph(str(r))

    doc.save(filename)






def rapport_text(agregation, analyse_jeur, analyse_Semaine, analyse_Mois, clients_10Tops, clients_Ayant_Non_Regles, top_resturants,
        top_heur, df_Anormalis, s_User, r, filename="documents_rapport/rapport.txt" 
        ):
    with open(filename, "w", encoding="utf-8") as f:
                f.write("***** RAPPORT JOURNALIER *****\n")

                f.write("\n======> Montants totaux par période Jour , Semaine , Mois : \n")
                f.write(f"{agregation.head()}\n")
                f.write(f"---- nomber des columns = {len(agregation)}\n")

                f.write("\n======> Moyenne des soldes des cartes prépayées : \n")
                f.write(f"{analyse_jeur.head()}\n")
                f.write(f"---- nomber des columns = {len(analyse_jeur)}\n")
                f.write(f"{analyse_Semaine.head()}\n")
                f.write(f"---- nomber des columns = {len(analyse_Semaine)}\n")
                f.write(f"{analyse_Mois.head()}\n")
                f.write(f"---- nomber des columns = {len(analyse_Mois)}\n")

                f.write("\n======> Clients les plus dépensiers : \n")
                f.write(f"{clients_10Tops.reset_index().head()}\n")
                f.write(f"---- nomber des columns = {len(clients_10Tops)}\n")
                f.write("\n======> Clients à risque : \n")
                f.write(f"{clients_Ayant_Non_Regles.reset_index().head()}\n")
                f.write(f"---- nomber des columns = {len(clients_Ayant_Non_Regles)}\n")

                f.write("\n======> Top restaurants de pointe : \n")
                f.write(f"{top_resturants.reset_index().head()}\n")
                f.write(f"---- nomber des columns = {len(top_resturants)}\n")
                f.write("\n======> Top heures de pointe : \n")
                f.write(f"{top_heur.reset_index().head()}\n")
                f.write(f"---- nomber des columns = {len(top_heur)}\n")

                f.write("\n======> Anomalies détectées : \n")
                f.write(f"{df_Anormalis.reset_index().head()}\n")
                f.write(f"---- nomber des columns = {len(df_Anormalis)}\n")

                f.write("\n======> Performance des Caissiers : \n")
                f.write(f"{s_User.reset_index().head()}\n")
                f.write(f"---- nomber des columns = {len(s_User)}\n")

                f.write("\n======> la Corrélations entre column Solde_CPP et Montant_Rgl : \n")
                f.write(f"{r}\n")
